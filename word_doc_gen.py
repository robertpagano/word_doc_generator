# %%
from docx import Document

import pandas as pd

import glob

from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.style import WD_STYLE_TYPE

from docxcompose.composer import Composer

from pathlib import Path

# %%
## This version extracts all needed data from filepaths list and doc list, directly from Flow API call

## Below is the final version of accessing files
## - It takes in 2 lists:
##  + filepaths_to_docs
##  + filepaths_to_summs
## - It will need to go inside of make_master_file, which is the file that will be used
## in the api

filepaths_to_docs = glob.glob('articles/**/*.docx')
filepaths_to_summs = glob.glob('articles/**/*.txt')

# %%

def make_doc_dataframe(filepaths_to_docs, filepaths_to_summs):

    '''
    this takes in two lists:
    filepaths_to_docs - filepaths to all of the articles in word doc form
    filepaths_to_summs - approved summaries of the articles, in txt form with "name + '_summary'" naming convention

    this function parses the paths in filepaths_to_docs and text in filepaths_to_summs to create a dataframe that
    is used to build the master document. this dataframe is sorted by section, and can be ordered in any section order
    by changing the value of 'section_order'.

    the columns in the dataframe are:
    
        - article_name - parsed article_name from filepath
        - section - parsed section content from filepath
        - doc_object - python docx object read in from filepath
        - summ_text - text extracted from each summary file
        - new_section - True if it's the first row of a specific section

    '''

    cols = ['article_name', 'section', 'doc_object', 'summ_text'] # new_section column created below

    temp_list = []

    # each list will be ['article_name', 'section', 'doc_object', 'summ_text'] for one path pair, 
    # then be appended to dataframe

    df = pd.DataFrame(columns=cols)

    # zip_dict = dict(zip(filepaths_list, zip(doc_list, summ_list_txt)))
    zipped_paths_dict = dict(zip(filepaths_to_docs, filepaths_to_summs))

    for doc_path, txt_path  in zipped_paths_dict.items():

        article_name = doc_path.split('\\')[-1][:-5]
        section = doc_path.split('\\')[1]
        article_doc = Document(doc_path)
        summ_text = Path(txt_path).read_text()

        temp_list.append(article_name)
        temp_list.append(section)
        temp_list.append(article_doc)
        temp_list.append(summ_text)

        # convert list to series using cols, then append series to dataframe
        temp_series = pd.Series(temp_list, index = df.columns)
        df = df.append(temp_series, ignore_index = True)
        temp_list = [] # empty the list for the next iteration

    # 'section_order' is the order we want the sections to be in the final document
    # this sorts the dataframe by a specific order in sections    
    section_order = ["Content + Training", "Product + Availability", "Programs + Offers", "Partner Update", "nocat"]
    df['section'] = pd.Categorical(df['section'], section_order)

    df = df.sort_values('section')

    # create new_section column to indicate when we need a section heading in master doc
    df['new_section'] = df['section'].shift().fillna('nocat') != df['section']

    return df

df = make_doc_dataframe(filepaths_to_docs, filepaths_to_summs)

df
# %%

def make_doc_dataframe_from_lists(filepaths_list, doc_list, summ_list_txt):
    '''
    new version - this will now take in a list of paths, and a list of files. DF will have the following fields:

    filename: filename ending with docx
    article_title: filename minus .docx
    section:
    month:
    doc_object: pulled from other list

    '''

    cols = ['filename', 'article_name', 'section', 'month', 'doc_object', 'summ_object', 'filepath']

    temp_list = []

    # each list will be [filename, article_name, section, month, doc_object, summ_object, filepath] for one path, 
    # then be appended to dataframe

    df = pd.DataFrame(columns=cols)

    zip_dict = dict(zip(filepaths_list, zip(doc_list, summ_list_txt)))

    # print(zip_dict)

    for path, docs in zip_dict.items():
        
        filename = path.split('/')[-1]
        article_name = filename[:-5]
        section = path.split('/')[-2]
        month = path.split('/')[-3]
        temp_list.append(filename)
        temp_list.append(article_name)
        temp_list.append(section)
        temp_list.append(month)
        temp_list.append(docs[0])
        temp_list.append(docs[1])
        temp_list.append(path)
        temp_series = pd.Series(temp_list, index = df.columns)
        df = df.append(temp_series, ignore_index = True)
        temp_list = []

    # # 'section_order' is the order we want the sections to be in the final document
    # # this sorts the dataframe by a specific order in sections    
    section_order = ["Content + Training", "Product + Availability", "Programs + Offers", "Partner Update", "nocat"]
    df['section'] = pd.Categorical(df['section'], section_order)

    df = df.sort_values('section')

    df['new_section'] = df['section'].shift().fillna('nocat') != df['section']
    # df['prev_section'] = df['prev_section'].fillna('nocat')
    # df['new_section'] = df['prev_section'] != df['section']

    
    return df

df = make_doc_dataframe_from_lists(filepaths_list, doc_list, summ_list_txt)

# %%
df.head(50)
# %%
'''
now I just need to make the document! look at word_doc_gen_v3 "make_master_file" etc.

for summary, try using actual code rather than calling API (but it may make sense to call API anyway)

- summary goes above full article, and call it abstract

  - if time, also code up the two different files

- clean up summaries - make sure to not grab title

- make logic to split up by section, and put divider images between the sections

  - try to make the divider images heading 1, and article names heading 2
  - make a document for each heading that lives in the repo, can pull from these using composer to add to the newsletter

- use font information that maria sent

  - All titles: Segoe UI Semibold font size 12 (black) 

    body/paragraphs - segoe ui font size 11 (black)   	

    unfilled bullets (like the photo - black) 

    Abstract/article : segoe ui semibold font size 11

- make sure in landscape mode (sideways)

- once all done, let's update flask api 2 repo, and azure repo

- header image (maria mentioned she had code for that, take a look at her old word doc file to see)
'''

# to test this, i'll need to make a list of real docs, and real summary docs
# so actually, we receive 3 lists

# %%
def create_doc(doc, summ_object, section, article_name, new_section = False):
    '''
    if new_section == True, then this will make a heading. This will only get triggered as True when a new section is
    seen in the dataframe

    '''
    
    # need to update to include section headers

    try:
        doc.styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)
        doc.styles.add_style('Heading 2', WD_STYLE_TYPE.PARAGRAPH)
    except:
        print('style already added')
        pass

    paragraphs = doc.paragraphs

    if new_section == True:

        p = paragraphs[0]
        section_paragraph = p.insert_paragraph_before(section)
        section_paragraph.style = doc.styles['Heading 1']
        title = p.insert_paragraph_before(article_name)
        title.style = doc.styles['Heading 2']
        abstract_title = p.insert_paragraph_before('Abstract:')
        abstract = p.insert_paragraph_before(summ_object)
        article_title = p.insert_paragraph_before('Article:')

        # TRY ADD HEADING INSTEAD OF THE STYLE STUFF SINCE BOTH HEADINGS ARE BEING ADDED AS NEW!!!!!!

    else:

        p = paragraphs[0]
        title = p.insert_paragraph_before(article_name)
        title.style = doc.styles['Heading 2']
        abstract_title = p.insert_paragraph_before('Abstract:')
        abstract = p.insert_paragraph_before(summ_object)
        article_title = p.insert_paragraph_before('Article:')
    
    # make a conditional where if article_name is in either paragraphs[0] of sumamry or article
    # delete it

    # if paragraphs[0].text == article_name:
    #     title = paragraphs[0]
    #     title.style = doc.styles['Heading 2']
    #     print(paragraphs[0].style)
    #     # print('same')
    # else:
    #     p = paragraphs[0]
    #     p.insert_paragraph_before(article_name)
    #     title = paragraphs[0]
    #     title.style = doc.styles['Heading 2']
        # print('not same')        
        # new_doc = Document()
        # doc.add_heading(article_name)
        # doc.add_paragraph(get_summ(doc_path))        
        # composer = Composer(new_doc)
        # composer.save(filename) 

    # if make_summary:
    #     summ_doc = Document()
    #     summ_doc.add_heading(article_name)
    #     summ_doc.add_paragraph(get_summ(doc))
    #     return summ_doc

    return doc

# %%

# test_doc = Document('Quick.docx')
# summ_object = 'fake summary'
# section = 'Partner Update'
# article_name = 'Quick'

# # %%
# new_doc = create_doc(test_doc, summ_object, section, article_name, new_section = True)

# # %%
# test_doc_2 = Document('Infrastructure update.docx')
# article_name_2 = 'Infrastructure update'

# new_doc_2 = create_doc(test_doc_2, summ_object, section, article_name_2, new_section = False)
# # %%
# new_doc.save('TESTINGTESTING.docx')
# # %%
# # %%
# master = new_doc
# composer = Composer(master)
# composer.append(test_doc_2)

# composer.save('testingcomposer.docx') 

# %%
# so the above doesnt seem to work on a single doc, but the other one didnt either
# hopefully it works when they all get smashed together, we will see!

def make_toc(doc):
    '''
    this function creates a table of contents object within a docx object, which will be called when the master files are created

    it indexes any text with "heading styles"

    if article "sections" are heading 1, and article titles are heading 2, it will take care of all of the proper formatting/indenting
    '''
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')  # creates a new element
    fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'   # change 1-3 depending on heading levels you need

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:t')
    fldChar3.text = "Right-click to update field."
    fldChar2.append(fldChar3)

    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')

    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar4)
    p_element = paragraph._p
    return doc
# %% 
def make_master_file(df, master_filename):

    ## THIS NEEDS TO TAKE TWO LISTS FROM API
    '''
    needs to take in the dictionary, parse out the variables, and change the calls
    to the functions to include those variables

    it uses the above functions to:
    
        - create the summaries for the text in each individual article
        - formats headings and makes a table of contents for each file

    '''
    toc = Document()
    paragraph = toc.add_paragraph('TABLE OF CONTENTS')
    toc = make_toc(toc)
    article_list = [toc]

    # for k, v in doc_dict.items():
    #     doc = v['doc']
    #     month = v['month']
    #     section = v['section']
    #     article_name = k
    #     article_list.append(create_doc(doc, month, section, article_name, summ))

    for index, row in df.iterrows():
        # first set all new section to FALSE, then test it and come back to add logic ot change this
        doc = row['doc_object']
        summ_object = row['summ_object']
        section = row['section']
        article_name = row['article_name']

        #check if article is first in new section, which would make us start a new section heading
        
        if row['new_section'] == True:
            article_list.append(create_doc(doc, summ_object, section, article_name, new_section = True))

        else:
            article_list.append(create_doc(doc, summ_object, section, article_name, new_section = False))

    master = article_list[0]
    composer = Composer(master)
    for document in article_list[1:]:
        composer.append(document)

    composer.save(master_filename) 

# %%
make_master_file(df, 'here_is_my_big_test.docx')
# %%
def make_masters(doc_dict, summ_filename, article_filename):
    '''
    this takes in the dictionary of documents and metadata

    it uses the above functions to:
    
        - create the summaries for the text in each individual article
        - formats headings and makes a table of contents for each file

    '''

    make_master_file(doc_dict, summ_filename, summ=True)
    make_master_file(doc_dict, article_filename, summ=False)

# make_masters(doc_dict, "master_summaries.docx", "master_articles.docx")

# %%
